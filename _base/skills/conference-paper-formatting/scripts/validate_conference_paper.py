#!/usr/bin/env python3
"""Validate a DOCX against the compact conference-paper baseline."""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError as exc:  # pragma: no cover
    raise SystemExit("python-docx is required: pip install python-docx") from exc


TARGET = {
    "page_width_mm": 187.0,
    "page_height_mm": 260.0,
    "top_mm": 25.0,
    "right_mm": 15.0,
    "bottom_mm": 20.0,
    "left_mm": 15.0,
    "column_gap_twips": 284,
}

INTRO_NAMES = {"introduction", "введение"}
REFERENCE_NAMES = {"references", "reference", "список литературы", "литература"}
FORBIDDEN_FRONT_HEADINGS = {
    "abstract",
    "аннотация",
    "keywords",
    "key words",
    "ключевые слова",
}
FIGURE_RE = re.compile(r"^(figure|fig\.|рисунок|рис\.)\s*\d+\s*[.:]", re.IGNORECASE)
TABLE_RE = re.compile(r"^(table|таблица)\s*\d+\s*[.:]", re.IGNORECASE)
REFERENCE_RE = re.compile(r"^\[(\d+)\]\s*")


class Findings:
    def __init__(self, strict_metadata: bool = False) -> None:
        self.failures: list[str] = []
        self.warnings: list[str] = []
        self.passes: list[str] = []
        self.strict_metadata = strict_metadata

    def passed(self, text: str) -> None:
        self.passes.append(text)

    def failed(self, text: str) -> None:
        self.failures.append(text)

    def warned(self, text: str) -> None:
        self.warnings.append(text)

    def metadata(self, text: str) -> None:
        (self.failed if self.strict_metadata else self.warned)(text)

    def print(self) -> None:
        for text in self.passes:
            print(f"[PASS] {text}")
        for text in self.warnings:
            print(f"[WARN] {text}")
        for text in self.failures:
            print(f"[FAIL] {text}")
        print(
            f"SUMMARY: {len(self.passes)} pass, "
            f"{len(self.warnings)} warning, {len(self.failures)} fail"
        )


def normalized_heading(text: str) -> str:
    value = " ".join(text.strip().lower().split())
    value = re.sub(r"^(?:[ivxlcdm]+|\d+)[.)]\s*", "", value)
    return value.rstrip(".:")


def close(actual: float, expected: float, tolerance: float = 1.0) -> bool:
    return abs(actual - expected) <= tolerance


def paragraph_has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing | .//w:pict"))


def iter_blocks(document: Document):
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield "table", Table(child, document)


def previous_meaningful(blocks, start: int):
    for index in range(start - 1, -1, -1):
        kind, obj = blocks[index]
        if kind == "table" or obj.text.strip() or paragraph_has_drawing(obj):
            return kind, obj
    return None


def next_meaningful(blocks, start: int):
    for index in range(start + 1, len(blocks)):
        kind, obj = blocks[index]
        if kind == "table" or obj.text.strip() or paragraph_has_drawing(obj):
            return kind, obj
    return None


def check_numbering_format(docx: Path, paragraph: Paragraph, findings: Findings) -> None:
    ppr = paragraph._p.pPr
    num_pr = ppr.numPr if ppr is not None else None
    if num_pr is None or num_pr.numId is None:
        findings.failed("Первый заголовок не имеет автоматической нумерации.")
        return
    num_id = num_pr.numId.val
    with ZipFile(docx) as archive:
        if "word/numbering.xml" not in archive.namelist():
            findings.failed("В DOCX отсутствует numbering.xml.")
            return
        root = ET.fromstring(archive.read("word/numbering.xml"))
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    abstract_id = None
    for node in root.findall("w:num", ns):
        if node.get(f"{{{ns['w']}}}numId") == str(num_id):
            ref = node.find("w:abstractNumId", ns)
            abstract_id = ref.get(f"{{{ns['w']}}}val") if ref is not None else None
            break
    fmt = None
    if abstract_id is not None:
        for node in root.findall("w:abstractNum", ns):
            if node.get(f"{{{ns['w']}}}abstractNumId") == abstract_id:
                level = node.find("w:lvl[@w:ilvl='0']", ns)
                num_fmt = level.find("w:numFmt", ns) if level is not None else None
                fmt = num_fmt.get(f"{{{ns['w']}}}val") if num_fmt is not None else None
                break
    if fmt == "upperRoman":
        findings.passed("Основные разделы используют римскую нумерацию.")
    else:
        findings.failed(f"Ожидалась римская нумерация разделов, найдено: {fmt!r}.")


def check_layout(document: Document, findings: Findings) -> None:
    if len(document.sections) != 1:
        findings.warned(f"Найдено разделов Word: {len(document.sections)}; базовый образец использует один.")
    layout_ok = True
    for index, section in enumerate(document.sections, 1):
        values = {
            "page_width_mm": section.page_width.mm,
            "page_height_mm": section.page_height.mm,
            "top_mm": section.top_margin.mm,
            "right_mm": section.right_margin.mm,
            "bottom_mm": section.bottom_margin.mm,
            "left_mm": section.left_margin.mm,
        }
        for key, expected in TARGET.items():
            if key == "column_gap_twips":
                continue
            if not close(values[key], expected):
                layout_ok = False
                findings.failed(
                    f"Раздел {index}: {key} = {values[key]:.2f}, ожидалось {expected:.2f}."
                )
        cols = section._sectPr.find(qn("w:cols"))
        count = int(cols.get(qn("w:num"), "1")) if cols is not None else 1
        gap = int(cols.get(qn("w:space"), "0")) if cols is not None else 0
        if count != 2:
            layout_ok = False
            findings.failed(f"Раздел {index}: ожидались две колонки, найдено {count}.")
        if abs(gap - TARGET["column_gap_twips"]) > 30:
            layout_ok = False
            findings.failed(
                f"Раздел {index}: промежуток колонок {gap} твипов, ожидалось около 284."
            )
    if layout_ok:
        findings.passed("Размер страницы, поля и две колонки соответствуют образцу.")


def check_font(document: Document, findings: Findings) -> None:
    inspected = 0
    matching = 0
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            inspected += 1
            name = run.font.name
            size = run.font.size.pt if run.font.size else None
            if name == "Times New Roman" and size is not None and abs(size - 10.0) <= 0.1:
                matching += 1
    ratio = matching / inspected if inspected else 0.0
    if ratio >= 0.85:
        findings.passed(f"Основной текст преимущественно Times New Roman 10 пт ({ratio:.0%}).")
    else:
        findings.warned(
            f"Только {ratio:.0%} явно оформленных фрагментов имеют Times New Roman 10 пт. "
            "Проверить стили и наследование шрифта."
        )


def check_tables(document: Document, findings: Findings) -> None:
    vertical = []
    for index, table in enumerate(document.tables, 1):
        borders = table._tbl.tblPr.find(qn("w:tblBorders"))
        if borders is None:
            continue
        for edge in ("left", "right", "insideV"):
            node = borders.find(qn(f"w:{edge}"))
            if node is not None and node.get(qn("w:val"), "nil") not in {"nil", "none"}:
                vertical.append((index, edge))
    if vertical:
        findings.failed(f"Обнаружены вертикальные границы таблиц: {vertical}.")
    else:
        findings.passed("Таблицы не используют вертикальные границы.")


def check_captions(document: Document, findings: Findings) -> None:
    blocks = list(iter_blocks(document))
    placement_warnings = []
    figure_labels = set()
    for index, (kind, obj) in enumerate(blocks):
        if kind != "paragraph":
            continue
        text = obj.text.strip()
        fig_match = FIGURE_RE.match(text)
        table_match = TABLE_RE.match(text)
        if fig_match:
            figure_labels.add(fig_match.group(1).lower())
            previous = previous_meaningful(blocks, index)
            has_adjacent_image = paragraph_has_drawing(obj) or (
                previous
                and previous[0] == "paragraph"
                and paragraph_has_drawing(previous[1])
            )
            if not has_adjacent_image:
                placement_warnings.append(f"подпись рисунка без изображения непосредственно перед ней: {text[:50]}")
            first_text_run = next((run for run in obj.runs if run.text.strip()), None)
            if first_text_run is not None and not first_text_run.bold:
                placement_warnings.append(f"метка рисунка не выделена полужирным: {text[:50]}")
        if table_match:
            following = next_meaningful(blocks, index)
            if not following or following[0] != "table":
                placement_warnings.append(f"подпись таблицы не находится непосредственно над таблицей: {text[:50]}")
            first_text_run = next((run for run in obj.runs if run.text.strip()), None)
            if first_text_run is not None and not first_text_run.bold:
                placement_warnings.append(f"метка таблицы не выделена полужирным: {text[:50]}")
    if len(figure_labels) > 1:
        placement_warnings.append(f"смешаны варианты метки рисунков: {sorted(figure_labels)}")
    if placement_warnings:
        for warning in placement_warnings:
            findings.warned(warning)
    else:
        findings.passed("Подписи рисунков и таблиц расположены и оформлены согласованно.")


def check_structure(document: Document, docx: Path, findings: Findings) -> None:
    visible = [(index, paragraph) for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip()]
    if not visible:
        findings.failed("Документ не содержит текста.")
        return
    _, first = visible[0]
    first_name = normalized_heading(first.text)
    if first_name in INTRO_NAMES:
        findings.passed("Документ начинается непосредственно с введения.")
        check_numbering_format(docx, first, findings)
    else:
        findings.failed(f"Первый непустой абзац должен быть введением, найдено: {first.text!r}.")

    forbidden = []
    for index, paragraph in visible:
        name = normalized_heading(paragraph.text)
        if name in FORBIDDEN_FRONT_HEADINGS:
            forbidden.append((index, paragraph.text))
    if forbidden:
        findings.failed(f"Обнаружены запрещённые служебные разделы: {forbidden}.")
    else:
        findings.passed("Титульный блок, аннотация и ключевые слова не обнаружены.")

    reference_candidates = [
        (index, paragraph)
        for index, paragraph in visible
        if normalized_heading(paragraph.text) in REFERENCE_NAMES
    ]
    if not reference_candidates:
        findings.failed("Не найден последний раздел со списком литературы.")
        return
    ref_index, ref_heading = reference_candidates[-1]
    ref_num_pr = ref_heading._p.pPr.numPr if ref_heading._p.pPr is not None else None
    ref_num_id = ref_num_pr.numId.val if ref_num_pr is not None and ref_num_pr.numId is not None else 0
    if ref_num_id != 0:
        findings.failed("Заголовок списка литературы не должен иметь номер раздела.")
    entries = []
    unexpected = []
    for index, paragraph in visible:
        if index <= ref_index:
            continue
        match = REFERENCE_RE.match(paragraph.text.strip())
        if match:
            entries.append(int(match.group(1)))
        else:
            unexpected.append((index, paragraph.text[:80]))
    if unexpected:
        findings.failed(f"После заголовка списка литературы найден иной текст: {unexpected}.")
    elif not entries:
        findings.failed("Список литературы не содержит нумерованных записей.")
    elif entries != list(range(1, len(entries) + 1)):
        findings.failed(f"Нумерация списка литературы нарушена: {entries}.")
    else:
        findings.passed("Список литературы является последним разделом и имеет сквозную нумерацию.")


def check_headers_and_metadata(document: Document, docx: Path, findings: Findings) -> None:
    visible_header_footer = []
    for index, section in enumerate(document.sections, 1):
        header_text = " ".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
        footer_text = " ".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
        if header_text or footer_text:
            visible_header_footer.append((index, header_text, footer_text))
    with ZipFile(docx) as archive:
        names = archive.namelist()
        field_text = b"\n".join(
            archive.read(name)
            for name in names
            if name.startswith("word/header") or name.startswith("word/footer")
        ).upper()
        has_page_field = b"PAGE" in field_text
        has_comments = any("comments" in name for name in names)
        document_root = ET.fromstring(archive.read("word/document.xml"))
        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        tracked = bool(
            document_root.findall(f".//{{{w_ns}}}ins")
            or document_root.findall(f".//{{{w_ns}}}del")
            or document_root.findall(f".//{{{w_ns}}}moveFrom")
            or document_root.findall(f".//{{{w_ns}}}moveTo")
        )
    if visible_header_footer or has_page_field:
        findings.failed("Обнаружены видимые колонтитулы или поле номера страницы.")
    else:
        findings.passed("Колонтитулы и номера страниц отсутствуют.")
    if has_comments or tracked:
        findings.failed("В документе остались комментарии или отслеживаемые исправления.")
    else:
        findings.passed("Комментарии и отслеживаемые исправления отсутствуют.")

    author = (document.core_properties.author or "").strip()
    modified_by = (document.core_properties.last_modified_by or "").strip()
    if author or modified_by:
        findings.metadata(
            f"В свойствах DOCX остались персональные данные: author={author!r}, lastModifiedBy={modified_by!r}."
        )
    else:
        findings.passed("Персональные поля author и lastModifiedBy очищены.")


def validate(docx: Path, strict_metadata: bool) -> int:
    findings = Findings(strict_metadata=strict_metadata)
    if not docx.is_file():
        print(f"[FAIL] File not found: {docx}")
        return 2
    try:
        with ZipFile(docx) as archive:
            bad = archive.testzip()
            if bad:
                findings.failed(f"Повреждён элемент ZIP-контейнера: {bad}.")
            else:
                findings.passed("ZIP-контейнер DOCX корректен.")
            for name in archive.namelist():
                if name.endswith(".xml") or name.endswith(".rels"):
                    ET.fromstring(archive.read(name))
        document = Document(docx)
    except Exception as exc:
        print(f"[FAIL] DOCX cannot be parsed: {exc}")
        return 2

    check_structure(document, docx, findings)
    check_layout(document, findings)
    check_font(document, findings)
    check_tables(document, findings)
    check_captions(document, findings)
    check_headers_and_metadata(document, docx, findings)
    findings.print()
    return 1 if findings.failures else 0


def main() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path)
    parser.add_argument(
        "--strict-metadata",
        action="store_true",
        help="Treat author and lastModifiedBy metadata as failures.",
    )
    args = parser.parse_args()
    raise SystemExit(validate(args.docx.resolve(), args.strict_metadata))


if __name__ == "__main__":
    main()
